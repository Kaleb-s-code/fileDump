'''
Created on Mar 13, 2019

The purpose of this module is to house all of the methods relating
to generating the spending report. 
@author: Kaleb
'''
from datetime import datetime
import dbMethods
from docx import Document
from docx.shared import Pt
from docx.shared import Inches


'''This method should generate a simple .txt file 
with a summary of the transactions for a given range'''
def generateGeneralReport():
    document = Document()
    
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(10)
    
    headingString = 'General Transaction Report'
    dateString = 'Generated on: ', str(datetime.now().strftime('%Y-%m-%d')), '.'
    numberOfTransactions = "Total number of transactions: {}".format(
        dbMethods.getTotalNumberOfTransactions(0, '', '', '', '', ''))
    totalPositive = "Total Positive: ${:,.2f}".format(dbMethods.getGeneralCashFlow(1))
    totalNegative = 'Total Negative: -${:,.2f}'.format(dbMethods.getGeneralCashFlow(0))
    document.add_heading(headingString, 0)
    document.add_heading(dateString)

    document.add_heading("Overview", level=1)
    
    document.add_paragraph(
        numberOfTransactions
    )
    document.add_paragraph(
        totalPositive
    )
    document.add_paragraph(
        totalNegative
    )
    records = (
        dbMethods.getTransactionsIntoArrayDescinding(0, '', '', '', '', '')
    )
    document.add_heading("Breakdown:", level=1)
    table = document.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item Id'
    hdr_cells[1].text = 'Date of Transaction'
    hdr_cells[2].text = 'Purchaser'
    hdr_cells[3].text = 'Vendor'
    hdr_cells[4].text = 'Description'
    hdr_cells[5].text = 'Category'
    hdr_cells[6].text = 'Amount'
    for itemId, dateOfTrans, purchaser, vendor, desc, cat, amount in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(itemId)
        row_cells[1].text = dateOfTrans
        row_cells[2].text = purchaser
        row_cells[3].text = vendor
        row_cells[4].text = desc
        row_cells[5].text = cat
        row_cells[6].text = amount

    document.save('GeneralReport{}.docx'.format(str(datetime.now().strftime('%Y-%m-%d'))))
    print('**Report Generated**')

    
def generateReportByDate(theBeginDate, theEndDate):
    document = Document()
    
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(10)
    
    headingString = 'Transactions by Date Report'
    dateString = 'Generated on: ', str(datetime.now().strftime('%Y-%m-%d')), '.'
    numberOfTransactions = "Total number of transactions: {}".format(dbMethods.getTotalNumberOfTransactions(1, '', '', '', theBeginDate, theEndDate))
    totalSpent = "Total spent in date range: ${:,.2f}".format(dbMethods.getTheTotalSpent(1,  '', '', '', theBeginDate, theEndDate))

    document.add_heading(headingString, 0)
    document.add_heading(dateString)
    
    document.add_heading("Overview", level=1)
    
    document.add_paragraph(
        numberOfTransactions
    )
    document.add_paragraph(
        totalSpent
    )
    records = (
        dbMethods.getTransactionsIntoArrayDescinding(1,  '', '', '', theBeginDate, theEndDate)
    )
    document.add_heading("Breakdown:", level=1)
    table = document.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item Id'
    hdr_cells[1].text = 'Date of Transaction'
    hdr_cells[2].text = 'Purchaser'
    hdr_cells[3].text = 'Vendor'
    hdr_cells[4].text = 'Description'
    hdr_cells[5].text = 'Category'
    hdr_cells[6].text = 'Amount'
    for itemId, dateOfTrans, purchaser, vendor, desc, cat, amount in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(itemId)
        row_cells[1].text = dateOfTrans
        row_cells[2].text = purchaser
        row_cells[3].text = vendor
        row_cells[4].text = desc
        row_cells[5].text = cat
        row_cells[6].text = amount

    document.save('SpendingByDate{}.docx'.format(str(datetime.now().strftime('%Y-%m-%d'))))
    print('**Report Generated**')

def generateReportByPurchaser(thePurchaser):
    document = Document()
    
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(10)
    
    headingString = 'Transactions by Purchaser Report'
    dateString = 'Generated on: ', str(datetime.now().strftime('%Y-%m-%d')), '| for: [', thePurchaser, ']'
    numberOfTransactions = "Total number of transactions: {}".format(dbMethods.getTotalNumberOfTransactions(2,  '', '', thePurchaser, '', ''))
    totalSpent = "Total spent by purchaser: ${:,.2f}".format(dbMethods.getTheTotalSpent(2, '', '', thePurchaser, '', ''))

    document.add_heading(headingString, 0)
    document.add_heading(dateString)
    
    
    document.add_heading("Overview", level=1)
    
    document.add_paragraph(
        numberOfTransactions
    )
    document.add_paragraph(
        totalSpent
    )
    records = (
        dbMethods.getTransactionsIntoArrayDescinding(2, '', '', thePurchaser, '', '')
    )
    document.add_heading("Breakdown:", level=1)
    table = document.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item Id'
    hdr_cells[1].text = 'Date of Transaction'
    hdr_cells[2].text = 'Purchaser'
    hdr_cells[3].text = 'Vendor'
    hdr_cells[4].text = 'Description'
    hdr_cells[5].text = 'Category'
    hdr_cells[6].text = 'Amount'
    for itemId, dateOfTrans, purchaser, vendor, desc, cat, amount in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(itemId)
        row_cells[1].text = dateOfTrans
        row_cells[2].text = purchaser
        row_cells[3].text = vendor
        row_cells[4].text = desc
        row_cells[5].text = cat
        row_cells[6].text = amount

    document.save('SpendingByPurchaser{}[{}].docx'.format(str(datetime.now().strftime('%Y-%m-%d')), thePurchaser))
    print('**Report Generated**')

def generateReportByAmount(theFile, theAmount):
    time = [datetime.now().strftime('%Y-%m-%d')]
    theFile.write('Transaction Summary:\n')
    theFile.write('Date: ' + str(time) + '\n')
    theFile.write('=' * 40 + '\n\n')
    theFile.write(str(dbMethods.viewTransactions(3, '', '', '', theAmount, '', '')))
    theFile.close()
    print('\n**Session Report Generated**')

def generateReportByVendor(theVendor):
    document = Document()
    
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(10)
    
    headingString = 'Transactions by Vendor Report'
    dateString = 'Generated on: ', str(datetime.now().strftime('%Y-%m-%d')), '.'
    numberOfTransactions = "Total number of transactions: {}".format(dbMethods.getTotalNumberOfTransactions(3, '', theVendor, '', '', ''))
    totalSpent = "Total spent by vendor: ${:,.2f}".format(dbMethods.getTheTotalSpent(3, '', theVendor, '', '', ''))

    document.add_heading(headingString, 0)
    document.add_heading(dateString)
    
    document.add_heading("Overview", level=1)
    
    document.add_paragraph(
        numberOfTransactions
    )
    document.add_paragraph(
        totalSpent
    )
    records = (
        dbMethods.getTransactionsIntoArrayDescinding(3, '', theVendor, '', '', '')
    )
    document.add_heading("Breakdown:", level=1)
    table = document.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item Id'
    hdr_cells[1].text = 'Date of Transaction'
    hdr_cells[2].text = 'Purchaser'
    hdr_cells[3].text = 'Vendor'
    hdr_cells[4].text = 'Description'
    hdr_cells[5].text = 'Category'
    hdr_cells[6].text = 'Amount'
    for itemId, dateOfTrans, purchaser, vendor, desc, cat, amount in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(itemId)
        row_cells[1].text = dateOfTrans
        row_cells[2].text = purchaser
        row_cells[3].text = vendor
        row_cells[4].text = desc
        row_cells[5].text = cat
        row_cells[6].text = amount

    document.save('SpendingByVendor{}.docx'.format(str(datetime.now().strftime('%Y-%m-%d'))))
    print('**Report Generated**')
    
    
def generateReportByCategory(theCat):
    document = Document()
    
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(10)
    
    headingString = 'Transactions by Category Report'
    dateString = 'Generated on: ', str(datetime.now().strftime('%Y-%m-%d')), '.'
    numberOfTransactions = "Total number of transactions: {}".format(dbMethods.getTotalNumberOfTransactions(4, theCat, '', '', '', ''))
    totalSpent = "Total spent in category: ${:,.2f}".format(dbMethods.getTheTotalSpent(4, theCat, '', '', '', ''))

    document.add_heading(headingString, 0)
    document.add_heading(dateString)
    
    document.add_heading("Overview", level=1)
    
    document.add_paragraph(
        numberOfTransactions
    )
    document.add_paragraph(
        totalSpent
    )
    records = (
        dbMethods.getTransactionsIntoArrayDescinding(4, theCat, '', '', '', '')
    )
    document.add_heading("Breakdown:", level=1)
    table = document.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item Id'
    hdr_cells[1].text = 'Date of Transaction'
    hdr_cells[2].text = 'Purchaser'
    hdr_cells[3].text = 'Vendor'
    hdr_cells[4].text = 'Description'
    hdr_cells[5].text = 'Category'
    hdr_cells[6].text = 'Amount'
    for itemId, dateOfTrans, purchaser, vendor, desc, cat, amount in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(itemId)
        row_cells[1].text = dateOfTrans
        row_cells[2].text = purchaser
        row_cells[3].text = vendor
        row_cells[4].text = desc
        row_cells[5].text = cat
        row_cells[6].text = amount

    document.save('SpendingByCategory{}.docx'.format(str(datetime.now().strftime('%Y-%m-%d'))))
    print('**Report Generated**')
    
def createWordDocumentOfBudget():
    document = Document()
    
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(10)
    
    totalBudgeted = "Total Amount For Budgeted Items: ${:,.2f} | Total Monthly Expenses: ${:,.2f}".format(
        dbMethods.getTotalBudgeted(), dbMethods.getTotalMonthlyExpenses())
    accounts = "Total In Accounts:\n{}".format(dbMethods.viewAccounts())
    amountNeeded = "Amount Needed To Cover Expenses (3611): {}".format(dbMethods.getAmountNeeded())
    totalSaved = "Total Period Savings: {}".format(dbMethods.getTotalSaved())
    expensesVsIncome = 'Monthly Expenses VS. Income: ${:,.2f}'.format(dbMethods.getTotalMonthlySaved())
    totalCash = "Total Cash Withdrawal: ${:,.2f}".format(dbMethods.getTotalCashWithdrawal())


    headingString = 'The Budget'
    dateString = 'Generated on: ', str(datetime.now().strftime('%Y-%m-%d')), '.'
    document.add_heading(headingString, 0)
    document.add_heading(dateString)
    
    document.add_heading("Financial Overview", level=1)
    
    document.add_paragraph(
        totalBudgeted
    )
    document.add_paragraph(
        accounts
    )
    document.add_paragraph (
        amountNeeded
    )
    document.add_paragraph (
        "{} | {}".format(totalSaved, expensesVsIncome)
    )
    document.add_paragraph (
        totalCash
    )
    records = (
        dbMethods.getDbItemsIntoArray()
    )
    document.add_heading("Breakdown:", level=1)
    table = document.add_table(rows=1, cols=8)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item Id'
    hdr_cells[0].width = Inches(0.5)
    hdr_cells[1].text = 'Date Last Paid'
    hdr_cells[1].width = Inches(2)
    hdr_cells[2].text = 'Item Name'
    hdr_cells[3].text = 'Current Value'
    hdr_cells[4].text = 'Budgeted Value'
    hdr_cells[5].text = 'Expected Monthly'
    hdr_cells[6].text = 'Due Date'
    hdr_cells[6].width = Inches(2)
    hdr_cells[7].text = 'Notes'
    for itemId, lastPaid, name, current, budgValue, exp, due, notes in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(itemId)
        row_cells[1].text = lastPaid
        row_cells[2].text = name
        row_cells[3].text = current
        row_cells[4].text = budgValue
        row_cells[5].text = exp
        row_cells[6].text = due
        row_cells[7].text = notes
        

    document.save('TheBudget{}.docx'.format(str(datetime.now().strftime('%Y-%m-%d'))))
    print('**Report Generated***')

